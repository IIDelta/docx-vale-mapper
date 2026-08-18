from __future__ import annotations
import os
from pathlib import Path
from docx import Document

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures" / "generated"

def ensure_fixtures_dir():
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

def create_fixture_a():
    doc = Document()
    doc.add_paragraph("Fixture A - Auto-Fix Candidate Document")
    doc.add_paragraph("This document contains auto-fix candidates.")
    
    # ordinary spaces between numbers and units
    doc.add_paragraph("The patient took 5 mg of the medication.")
    doc.add_paragraph("Treatment lasted for 3 days without issues.")
    doc.add_paragraph("Dosage was 0.75 mg/kg initially.")
    
    # invalid math spacing
    doc.add_paragraph("The value was < 4 in all cases.")
    doc.add_paragraph("Age ≥ 20 years is required.")
    doc.add_paragraph("The margin of error was ± 2 percent.")
    
    # label capitalization candidate
    doc.add_paragraph("Please refer to version 5 of the protocol.")
    
    # repeated auto-fix candidates in separate paragraphs
    doc.add_paragraph("First instance of 5 mg.")
    doc.add_paragraph("Second instance of 5 mg.")
    
    # repeated auto-fix candidates in the same paragraph
    doc.add_paragraph("Here we have 5 mg and another 5 mg in the same sentence.")
    
    doc.save(FIXTURES_DIR / "Fixture_A_AutoFix.docx")

def create_fixture_b():
    doc = Document()
    doc.add_paragraph("Fixture B - Comment Aggregation Document")
    
    # 3 occurrences of a comment rule
    doc.add_paragraph("This is RuleX occurrence one.")
    doc.add_paragraph("This is RuleX occurrence two.")
    doc.add_paragraph("This is RuleX occurrence three.")
    
    # 5 occurrences of a comment rule (RuleY)
    doc.add_paragraph("This is RuleY occurrence one.")
    doc.add_paragraph("This is RuleY occurrence two.")
    doc.add_paragraph("This is RuleY occurrence three.")
    doc.add_paragraph("This is RuleY occurrence four.")
    doc.add_paragraph("This is RuleY occurrence five.")
    
    # EndOfTrial-style repeated finding
    doc.add_paragraph("End of trial finding one.")
    doc.add_paragraph("End of trial finding two.")
    doc.add_paragraph("End of trial finding three.")
    doc.add_paragraph("End of trial finding four.")
    doc.add_paragraph("End of trial finding five.")
    
    # TrialIntervention-style repeated finding
    doc.add_paragraph("Trial intervention one.")
    doc.add_paragraph("Trial intervention two.")
    doc.add_paragraph("Trial intervention three.")
    doc.add_paragraph("Trial intervention four.")
    doc.add_paragraph("Trial intervention five.")
    doc.add_paragraph("Trial intervention six.")
    
    doc.save(FIXTURES_DIR / "Fixture_B_Aggregation.docx")

def create_fixture_c():
    doc = Document()
    doc.add_paragraph("Fixture C - Ineligible Context Document")
    
    p = doc.add_paragraph("This is title page text.")
    p.style = "Title"
    
    p = doc.add_paragraph("Protocol summary content.")
    p.style = "Heading 1"
    
    table = doc.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    cell.text = "This is a table cell with 5 mg."
    
    p = doc.add_paragraph("Table 1. Table footnote text here.")
    p.style = "Footnote Text" # assuming style name
    
    p = doc.add_paragraph("Reference 1. Bibliography-like styled text.")
    p.style = "Bibliography" # assuming style name
    
    p = doc.add_paragraph("Reference-like styled text.")
    p.style = "List Paragraph"
    
    # Normal eligible body narrative
    doc.add_paragraph("This is a normal eligible body narrative text containing 5 mg.")
    
    doc.save(FIXTURES_DIR / "Fixture_C_Ineligible.docx")

def create_fixture_d():
    doc = Document()
    doc.add_paragraph("Fixture D - Output Integrity Document")
    
    p = doc.add_paragraph("This paragraph has an existing non-MVA comment.")
    # python-docx does not natively support adding comments easily, but we'll try or mock it
    # We might need to skip inserting real OOXML comments if python-docx doesn't support p.add_comment natively.
    # Actually python-docx doesn't have `add_comment`.
    # Let's just create a normal document. We'll verify MVA comments.
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    
    # Mixed formatting
    p = doc.add_paragraph()
    p.add_run("Here is some ").bold = True
    p.add_run("mixed ")
    p.add_run("formatting ").italic = True
    r = p.add_run("and superscript.")
    r.font.superscript = True
    
    # 5 mg (for NBSP text)
    doc.add_paragraph("Text with 5 mg to be replaced with NBSP.")
    
    doc.save(FIXTURES_DIR / "Fixture_D_Integrity.docx")

if __name__ == "__main__":
    ensure_fixtures_dir()
    create_fixture_a()
    create_fixture_b()
    create_fixture_c()
    create_fixture_d()
    print(f"Phase 2 synthetic fixtures successfully generated at: {FIXTURES_DIR}")
