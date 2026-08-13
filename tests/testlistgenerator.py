from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from abbreviations.legacyimport import (
    initialize_database,
    open_database,
)
from abbreviations.listgenerator import (
    build_generation_plan,
    generate_list_document,
)
from abbreviations.reviewpromote import (
    ensure_review_schema,
    upsert_registry_entry,
)


class ListGeneratorTests(unittest.TestCase):
    """Tests for B4.3 list-generation behavior."""

    def create_database(self) -> Path:
        """Create a temporary reviewed registry."""

        temporary_directory = tempfile.TemporaryDirectory()
        self.addCleanup(temporary_directory.cleanup)

        database_path = (
            Path(temporary_directory.name)
            / "abbreviations.sqlite"
        )

        initialize_database(database_path)

        with open_database(database_path) as connection:
            ensure_review_schema(connection)

            upsert_registry_entry(
                connection=connection,
                token="9vHPV",
                definition=(
                    "9-valent human papillomavirus"
                ),
                status="approved_expand",
                source_reference="test_source",
                replacement_token="",
                notes="Approved.",
            )

            upsert_registry_entry(
                connection=connection,
                token="AUC0-24",
                definition=(
                    "Area under the plasma concentration-time "
                    "curve over 24 hours"
                ),
                status="approved_expand",
                source_reference="test_source",
                replacement_token="",
                notes="Approved.",
            )

            upsert_registry_entry(
                connection=connection,
                token="CFR",
                definition="Code of Federal Regulations",
                status="approved_expand",
                source_reference="test_source",
                replacement_token="",
                notes="Approved.",
            )

            upsert_registry_entry(
                connection=connection,
                token="FDA",
                definition="",
                status="approved_no_expand",
                source_reference="test_source",
                replacement_token="",
                notes="Protected.",
            )

            upsert_registry_entry(
                connection=connection,
                token="FAS",
                definition="",
                status="ambiguous",
                source_reference="test_source",
                replacement_token="",
                notes="Ambiguous.",
            )

        return database_path

    def test_plan_sorts_and_excludes_protected_terms(
        self,
    ) -> None:
        database_path = self.create_database()

        candidate_report = {
            "candidates": [
                {"token": "CFR"},
                {"token": "FDA"},
                {"token": "AUC0-24"},
                {"token": "9vHPV"},
            ]
        }

        plan = build_generation_plan(
            database_path=database_path,
            candidate_report=candidate_report,
        )

        self.assertTrue(plan.can_generate)

        self.assertEqual(
            [entry.token for entry in plan.entries],
            [
                "9vHPV",
                "AUC0-24",
                "CFR",
            ],
        )

        self.assertEqual(
            plan.excluded_tokens,
            ["FDA"],
        )

        self.assertEqual(plan.blockers, [])

    def test_ambiguous_term_blocks_generation(self) -> None:
        database_path = self.create_database()

        candidate_report = {
            "candidates": [
                {"token": "CFR"},
                {"token": "FAS"},
            ]
        }

        plan = build_generation_plan(
            database_path=database_path,
            candidate_report=candidate_report,
        )

        self.assertFalse(plan.can_generate)

        self.assertEqual(
            len(plan.blockers),
            1,
        )

        self.assertEqual(
            plan.blockers[0].token,
            "FAS",
        )

        self.assertEqual(
            plan.blockers[0].status,
            "ambiguous",
        )

    def test_generated_docx_contains_sorted_rows(self) -> None:
        database_path = self.create_database()

        candidate_report = {
            "candidates": [
                {"token": "CFR"},
                {"token": "AUC0-24"},
                {"token": "9vHPV"},
            ]
        }

        plan = build_generation_plan(
            database_path=database_path,
            candidate_report=candidate_report,
        )

        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = (
                Path(temporary_directory)
                / "generated.docx"
            )

            generate_list_document(
                plan=plan,
                output_path=output_path,
            )

            document = Document(output_path)

        self.assertEqual(
            document.paragraphs[0].text,
            "LIST OF ABBREVIATIONS AND DEFINITION OF TERMS",
        )

        table = document.tables[0]

        self.assertEqual(
            table.cell(0, 0).text,
            "Abbreviation",
        )

        self.assertEqual(
            table.cell(0, 1).text,
            "Definition",
        )

        self.assertEqual(
            table.cell(1, 0).text,
            "9vHPV",
        )

        self.assertEqual(
            table.cell(2, 0).text,
            "AUC0-24",
        )

        self.assertEqual(
            table.cell(3, 0).text,
            "CFR",
        )


if __name__ == "__main__":
    unittest.main()
