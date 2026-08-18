import os
import shutil
from pathlib import Path
from docx import Document

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures" / "generated" / "v1"

def ensure_fixtures_dir():
    if FIXTURES_DIR.exists():
        shutil.rmtree(FIXTURES_DIR)
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

def create_base_docx(filename: str, content_func) -> Path:
    doc = Document()
    content_func(doc)
    out_path = FIXTURES_DIR / filename
    doc.save(out_path)
    return out_path

# 1. Rules and reports fixture (Report-only rules should not create comments)
def fixture_1(doc):
    doc.add_heading("Rules and Reports Fixture", level=1)
    # Trigger report_only rules: Clinical.DashSpacing and Clinical.MultiplePunctuationSpaces
    doc.add_paragraph("This paragraph has an en-dash - with spaces around it.")
    doc.add_paragraph("This paragraph has  multiple spaces after punctuation.")

# 2. Auto-fix fixture
def fixture_2(doc):
    doc.add_heading("Auto-fix Fixture", level=1)
    doc.add_paragraph("Here is an instance of 5 mg to fix.")
    doc.add_paragraph("A math operator with bad spacing < 4 is here.")
    doc.add_paragraph("Dotted abbreviations like i.e. need fixing.")

# 3. Aggregated comment fixture
def fixture_3(doc):
    doc.add_heading("Aggregated Comment Fixture", level=1)
    # Create 5 occurrences of the same rule to trigger aggregation
    # We will use "In order to" which is wordy
    for i in range(6):
        doc.add_paragraph("In order to test aggregation, we repeat this wordy phrase.")

# 4. Ineligible context fixture
def fixture_4(doc):
    doc.add_heading("Ineligible Context Fixture", level=1)
    p = doc.add_paragraph("Title Page Text with 5 mg")
    p.style = "Title"

# 5. Protected field fixture
def fixture_5(doc):
    doc.add_heading("Protected Field Fixture", level=1)
    doc.add_paragraph("A paragraph that will receive a form field with 5 mg.")

# 6. Existing comment preservation fixture
def fixture_6(doc):
    doc.add_heading("Existing Comment Preservation", level=1)
    doc.add_paragraph("A paragraph that has an existing comment from a reviewer. 5 mg is here.")

# 7. Table and footnote fixture
def fixture_7(doc):
    from docx.enum.style import WD_STYLE_TYPE
    if 'Footnote Text' not in doc.styles:
        doc.styles.add_style('Footnote Text', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_heading("Table and Footnote Fixture", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Table text with 5 mg inside."
    # We add a paragraph formatted as footnote text
    p = doc.add_paragraph("This is footnote text with 5 mg.")
    p.style = "Footnote Text"

# 8. Figure/caption fixture
def fixture_8(doc):
    from docx.enum.style import WD_STYLE_TYPE
    if 'Caption' not in doc.styles:
        doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    doc.add_heading("Figure and Caption Fixture", level=1)
    p = doc.add_paragraph("Figure 1. A caption with 5 mg.")
    p.style = "Caption"

# 9. Appendix fixture
def fixture_9(doc):
    doc.add_heading("Appendix Fixture", level=1)
    p = doc.add_paragraph("Appendix A. Supplementary data with 5 mg.")
    p.style = "Heading 1" 

# 10. Output integrity fixture
def fixture_10(doc):
    doc.add_heading("Output Integrity Fixture", level=1)
    p = doc.add_paragraph()
    p.add_run("bold ").bold = True
    p.add_run("italic ").italic = True
    p.add_run("and normal 5 mg.")

def post_process_with_com(path5: Path, path6: Path):
    # COM is unavailable on Linux. To fully test on Windows, this function could add
    # a real FORMTEXT field to Fixture 5 and a real comment to Fixture 6.
    pass

def main():
    ensure_fixtures_dir()
    
    paths = []
    paths.append(create_base_docx("Fixture_1_RulesReports.docx", fixture_1))
    paths.append(create_base_docx("Fixture_2_AutoFix.docx", fixture_2))
    paths.append(create_base_docx("Fixture_3_Aggregation.docx", fixture_3))
    paths.append(create_base_docx("Fixture_4_Ineligible.docx", fixture_4))
    paths.append(create_base_docx("Fixture_5_ProtectedField.docx", fixture_5))
    paths.append(create_base_docx("Fixture_6_ExistingComment.docx", fixture_6))
    paths.append(create_base_docx("Fixture_7_TableFootnote.docx", fixture_7))
    paths.append(create_base_docx("Fixture_8_FigureCaption.docx", fixture_8))
    paths.append(create_base_docx("Fixture_9_Appendix.docx", fixture_9))
    paths.append(create_base_docx("Fixture_10_OutputIntegrity.docx", fixture_10))
    
    post_process_with_com(paths[4], paths[5])
    print(f"Generated {len(paths)} V1 synthetic fixtures at: {FIXTURES_DIR}")

if __name__ == "__main__":
    main()
