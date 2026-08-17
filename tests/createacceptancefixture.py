from __future__ import annotations

from pathlib import Path

from docx import Document


OUTPUT_PATH = Path(__file__).resolve().parent / "fixtures" / "reportsonlyfixture.docx"


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Study Results For The Trial", level=1)
    document.add_paragraph(
        "Patients were randomized 1/1 and received 5 mg for 3 days."
    )
    document.add_paragraph(
        "The trial met its primary endpoint. (Smith et al, 2025)"
    )
    document.add_paragraph(
        "Escherichia coli was evaluated in vitro."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Population"
    table.cell(0, 1).text = "Total"
    table.cell(1, 0).text = "Safety population"
    table.cell(1, 1).text = "NA"
    document.save(OUTPUT_PATH)
    print(f"Reports-only acceptance fixture created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
