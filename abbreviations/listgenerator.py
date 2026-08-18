"""
Purpose:
    Manage abbreviation discovery, registry data, candidate review,
    and policy promotion.

Inputs:
    Candidate terms, registry data, review decisions, and policy files.

Outputs:
    Review reports, registry records, and effective abbreviation policy.

Must not:
    Insert Word comments.
    Modify Word documents directly.
"""

from __future__ import annotations

import argparse
import json
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document

from abbreviations.registryapi import resolve_token
from validators.abbreviationvalidator import (
    abbreviation_sort_key,
)


INCLUDED_STATUSES = {
    "approved_expand",
    "approved_list_only",
}

EXCLUDED_STATUSES = {
    "approved_no_expand",
    "ignored",
}

BLOCKING_STATUSES = {
    "unknown",
    "reviewed_candidate",
    "ambiguous",
    "deprecated",
}


@dataclass(frozen=True)
class ListEntry:
    """One abbreviation-definition row for generated output."""

    token: str
    definition: str
    status: str
    source_reference: str


@dataclass(frozen=True)
class GenerationBlocker:
    """One unresolved candidate preventing safe list generation."""

    token: str
    status: str
    message: str


@dataclass
class ListGenerationPlan:
    """Complete generation plan before any DOCX is created."""

    entries: list[ListEntry]
    excluded_tokens: list[str]
    blockers: list[GenerationBlocker]

    @property
    def can_generate(self) -> bool:
        """Return True only when no unresolved candidates remain."""

        return bool(self.entries) and not self.blockers

    def to_dict(self) -> dict[str, Any]:
        """Convert the generation plan into JSON-safe output."""

        return {
            "can_generate": self.can_generate,
            "entry_count": len(self.entries),
            "excluded_count": len(self.excluded_tokens),
            "blocker_count": len(self.blockers),
            "entries": [
                asdict(entry)
                for entry in self.entries
            ],
            "excluded_tokens": self.excluded_tokens,
            "blockers": [
                asdict(blocker)
                for blocker in self.blockers
            ],
        }


def load_candidate_report(
    report_path: Path,
) -> dict[str, Any]:
    """Load candidate report JSON."""

    with report_path.open(
        encoding="utf-8",
    ) as input_file:
        return json.load(input_file)


def blocker_message(
    token: str,
    status: str,
    replacement_token: str,
) -> str:
    """Return user-facing reason that a candidate blocks generation."""

    if status == "deprecated":
        if replacement_token:
            return (
                f"'{token}' is deprecated. "
                f"Use '{replacement_token}' before generating the list."
            )

        return (
            f"'{token}' is deprecated and requires review."
        )

    if status == "ambiguous":
        return (
            f"'{token}' has multiple possible definitions. "
            "Choose an approved definition first."
        )

    if status == "reviewed_candidate":
        return (
            f"'{token}' is reviewed but not approved "
            "for list generation."
        )

    return (
        f"'{token}' has no approved registry decision."
    )


def build_generation_plan(
    database_path: Path,
    candidate_report: dict[str, Any],
) -> ListGenerationPlan:
    """
    Build a compliant abbreviation-list plan.

    The current reviewed registry is authoritative. Candidate-report
    status values are not trusted as final policy decisions.
    """

    entries: list[ListEntry] = []
    excluded_tokens: list[str] = []
    blockers: list[GenerationBlocker] = []

    seen_tokens: set[str] = set()

    for candidate in candidate_report.get("candidates", []):
        token = str(
            candidate.get("token", "")
        ).strip()

        if not token:
            continue

        normalized_token = token.casefold()

        if normalized_token in seen_tokens:
            continue

        seen_tokens.add(normalized_token)

        resolution = resolve_token(
            database_path=database_path,
            token=token,
        )

        status = resolution.status

        if status in EXCLUDED_STATUSES:
            excluded_tokens.append(
                resolution.token or token
            )
            continue

        if status in INCLUDED_STATUSES:
            if not resolution.preferred_definition:
                blockers.append(
                    GenerationBlocker(
                        token=resolution.token or token,
                        status=status,
                        message=(
                            f"'{resolution.token or token}' is "
                            "approved for list use but has no "
                            "preferred definition."
                        ),
                    )
                )
                continue

            entries.append(
                ListEntry(
                    token=resolution.token or token,
                    definition=resolution.preferred_definition,
                    status=status,
                    source_reference=resolution.source_reference,
                )
            )
            continue

        if status in BLOCKING_STATUSES:
            blockers.append(
                GenerationBlocker(
                    token=resolution.token or token,
                    status=status,
                    message=blocker_message(
                        token=resolution.token or token,
                        status=status,
                        replacement_token=(
                            resolution.replacement_token
                        ),
                    ),
                )
            )
            continue

        blockers.append(
            GenerationBlocker(
                token=resolution.token or token,
                status=status,
                message=(
                    f"'{resolution.token or token}' has an "
                    "unsupported registry status."
                ),
            )
        )

    entries = sorted(
        entries,
        key=lambda entry: abbreviation_sort_key(
            entry.token
        ),
    )

    excluded_tokens = sorted(
        set(excluded_tokens),
        key=abbreviation_sort_key,
    )

    blockers = sorted(
        blockers,
        key=lambda blocker: abbreviation_sort_key(
            blocker.token
        ),
    )

    return ListGenerationPlan(
        entries=entries,
        excluded_tokens=excluded_tokens,
        blockers=blockers,
    )


def generate_list_document(
    plan: ListGenerationPlan,
    output_path: Path,
) -> None:
    """
    Generate a standalone List of Abbreviations DOCX.

    Generation is blocked if unresolved candidates remain.
    """

    if not plan.can_generate:
        raise ValueError(
            "List generation is blocked because unresolved "
            "candidates remain."
        )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    document.add_heading(
        "LIST OF ABBREVIATIONS AND DEFINITION OF TERMS",
        level=1,
    )

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    header_cells[0].text = "Abbreviation"
    header_cells[1].text = "Definition"

    for entry in plan.entries:
        cells = table.add_row().cells

        cells[0].text = entry.token
        cells[1].text = entry.definition

    document.save(output_path)


def write_generation_report(
    plan: ListGenerationPlan,
    output_path: Path,
) -> None:
    """Write a JSON generation-plan report."""

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    with output_path.open(
        "w",
        encoding="utf-8",
    ) as output_file:
        json.dump(
            plan.to_dict(),
            output_file,
            indent=2,
            ensure_ascii=False,
        )
        output_file.write("\n")


def parse_arguments() -> argparse.Namespace:
    """Parse list-generation command options."""

    parser = argparse.ArgumentParser(
        description=(
            "Generate a reviewed abbreviation-list DOCX from "
            "a candidate report and local registry."
        )
    )

    parser.add_argument(
        "--input",
        type=Path,
        required=True,
        help="Candidate report JSON path.",
    )

    parser.add_argument(
        "--database",
        type=Path,
        default=Path("data") / "abbreviations.sqlite",
        help=(
            "SQLite database path. "
            "Default: data/abbreviations.sqlite"
        ),
    )

    parser.add_argument(
        "--output",
        type=Path,
        required=True,
        help="Generated DOCX output path.",
    )

    parser.add_argument(
        "--report",
        type=Path,
        default=Path("reports") / "listgenerationreport.json",
        help=(
            "Generation-plan JSON report path. "
            "Default: reports/listgenerationreport.json"
        ),
    )

    return parser.parse_args()


def main() -> int:
    """Build plan, write report, and generate a DOCX if safe."""

    arguments = parse_arguments()

    if not arguments.input.is_file():
        print(
            "LIST GENERATION FAILED: "
            f"Candidate report not found: {arguments.input}"
        )
        return 2

    try:
        candidate_report = load_candidate_report(
            arguments.input
        )

        plan = build_generation_plan(
            database_path=arguments.database,
            candidate_report=candidate_report,
        )

        write_generation_report(
            plan=plan,
            output_path=arguments.report,
        )

        if not plan.can_generate:
            print(
                "LIST GENERATION BLOCKED: "
                f"{len(plan.blockers)} unresolved candidate(s)."
            )

            for blocker in plan.blockers:
                print(
                    f"- {blocker.token}: {blocker.message}"
                )

            print(
                f"Plan report: {arguments.report.resolve()}"
            )

            return 1

        generate_list_document(
            plan=plan,
            output_path=arguments.output,
        )

    except (
        OSError,
        ValueError,
        json.JSONDecodeError,
    ) as error:
        print(f"LIST GENERATION FAILED: {error}")
        return 2

    print(
        f"List entries generated: {len(plan.entries)}"
    )

    print(
        f"Excluded tokens: {len(plan.excluded_tokens)}"
    )

    print(
        f"DOCX: {arguments.output.resolve()}"
    )

    print(
        f"Plan report: {arguments.report.resolve()}"
    )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
